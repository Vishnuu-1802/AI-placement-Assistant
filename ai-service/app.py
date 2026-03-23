from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import pickle
import pandas as pd
import spacy
from serpapi import GoogleSearch
from dotenv import load_dotenv, find_dotenv
import os
import requests
import json
import re
from typing import Any, Dict, List
import time
import random
from collections import Counter
from io import BytesIO
from datetime import datetime

try:
    from pypdf import PdfReader
except Exception:
    try:
        from PyPDF2 import PdfReader
    except Exception:
        PdfReader = None

try:
    from docx import Document
except Exception:
    Document = None

# ------------------------------ 
# Flask Setup
# ------------------------------
app = Flask(__name__)
CORS(
    app,
    resources={r"/*": {"origins": ["http://localhost:5173", "http://localhost:5174", "http://127.0.0.1:5173", "http://127.0.0.1:5174"]}},
    supports_credentials=False,
)
TEST_CACHE: Dict[str, Dict[str, Any]] = {}
COMMON_TECH_SKILLS = {
    "python", "java", "javascript", "typescript", "react", "node", "node.js",
    "express", "flask", "django", "html", "css", "sql", "mongodb", "mysql",
    "postgresql", "aws", "azure", "gcp", "docker", "kubernetes", "git",
    "linux", "pandas", "numpy", "tensorflow", "pytorch", "machine learning",
    "devops", "ci/cd", "rest api", "graphql", "redis", "spark", "hadoop",
}
ACTION_VERBS = {
    "built", "developed", "implemented", "designed", "optimized", "deployed",
    "created", "improved", "led", "automated", "reduced", "increased", "managed",
    "integrated", "delivered"
}
NON_SKILL_TERMS = {
    "english", "hindi", "tamil", "telugu", "cgpa", "gpa", "percentage",
    "btech", "mtech", "b.e", "b.e.", "cse", "ece", "mech", "civil",
}

# Load environment variables
load_dotenv(find_dotenv())
SERP_API_KEY = os.getenv("SERP_API_KEY")
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")

print("\n==============================")
print("🔐 ENVIRONMENT CHECK")
print("SERP API key loaded:", bool(SERP_API_KEY))
print("DEEPSEEK API key loaded:", bool(DEEPSEEK_API_KEY))
print("==============================\n")

# ------------------------------ 
# Load NLP model
# ------------------------------
nlp = spacy.load("en_core_web_sm")

# ------------------------------ 
# Load ML model + encoders
# ------------------------------
model = pickle.load(open("career_model.pkl", "rb"))
encoder = pickle.load(open("skill_encoder.pkl", "rb"))

# ------------------------------ 
# Load dataset
# ------------------------------
df = pd.read_csv("synthetic_career_dataset_2000_extended.csv")


def _normalize_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip())


def _is_valid_skill_token(token: str) -> bool:
    t = (token or "").strip().lower()
    if len(t) < 2:
        return False
    if t in NON_SKILL_TERMS:
        return False

    # Exclude pure numeric/date/percentage-like tokens.
    if not re.search(r"[a-zA-Z]", t):
        return False
    if re.fullmatch(r"[\d\s./:%\-+]+", t):
        return False
    if re.fullmatch(r"\d{2,4}[-/]\d{1,4}", t):
        return False

    # Keep tech-like tokens and normal words.
    return True


def _extract_skills_from_text(text: str) -> List[str]:
    lowered = text.lower()
    detected = set()

    doc = nlp(text)
    for ent in doc.ents:
        token = ent.text.lower().strip()
        if len(token) > 1:
            detected.add(token)

    known_skills = set(map(str.lower, list(encoder.classes_))) | COMMON_TECH_SKILLS
    for skill in known_skills:
        if len(skill) > 1 and skill in lowered:
            detected.add(skill)

    return sorted({s.strip() for s in detected if _is_valid_skill_token(s)})


def _infer_section_presence(text: str) -> Dict[str, bool]:
    lowered = text.lower()
    patterns = {
        "summary": ["summary", "profile", "about me", "objective"],
        "skills": ["skills", "technical skills", "tooling", "technologies"],
        "projects": ["projects", "project experience", "portfolio"],
        "experience": ["experience", "work history", "employment"],
        "education": ["education", "academic", "university", "college"],
        "certifications": ["certifications", "certificates", "license"],
        "achievements": ["achievements", "awards", "accomplishments"],
    }
    return {k: any(token in lowered for token in v) for k, v in patterns.items()}


def _calculate_resume_analysis(text: str, skills: List[str]) -> Dict[str, Any]:
    words = text.split()
    word_count = len(words)
    text_lower = text.lower()
    sections = _infer_section_presence(text)
    section_count = sum(1 for value in sections.values() if value)

    action_verb_hits = sum(text_lower.count(v) for v in ACTION_VERBS)
    action_verb_score = min(100, action_verb_hits * 10)

    skill_depth = min(100, int(len(skills) * 6.5))
    detail_depth = min(100, int((word_count / 420) * 100))
    structure_score = min(100, int((section_count / len(sections)) * 100))

    final_score = int(
        0.38 * skill_depth
        + 0.32 * detail_depth
        + 0.20 * structure_score
        + 0.10 * action_verb_score
    )
    final_score = max(8, min(100, final_score))

    if final_score >= 80:
        quality_band = "Excellent"
    elif final_score >= 65:
        quality_band = "Strong"
    elif final_score >= 45:
        quality_band = "Developing"
    else:
        quality_band = "Needs Improvement"

    strengths: List[str] = []
    weaknesses: List[str] = []

    if len(skills) >= 8:
        strengths.append("Good spread of technical skills detected.")
    else:
        weaknesses.append("Technical skill coverage is limited for competitive roles.")

    if word_count >= 220:
        strengths.append("Resume has sufficient descriptive detail.")
    else:
        weaknesses.append("Resume is brief; add more quantified project and impact detail.")

    if section_count >= 5:
        strengths.append("Core resume sections are mostly present.")
    else:
        weaknesses.append("Missing key sections such as projects, experience, or certifications.")

    if action_verb_hits >= 4:
        strengths.append("Uses action-oriented language that improves impact.")
    else:
        weaknesses.append("Add stronger action verbs to communicate ownership and results.")

    target_role = None
    role_skills: List[str] = []
    missing_skills: List[str] = []
    fit = 0

    valid_skills_for_model = [s for s in skills if s in encoder.classes_]
    if valid_skills_for_model:
        try:
            vector = encoder.transform([valid_skills_for_model])
            target_role = str(model.predict(vector)[0])
            role_row = df[df["job_role"] == target_role].iloc[0]
            role_skills = [s.strip().lower() for s in str(role_row["skills"]).split(",") if s.strip()]
            missing_skills = sorted(list(set(role_skills) - set(valid_skills_for_model)))[:8]
            fit = int((len(set(role_skills) & set(valid_skills_for_model)) / max(1, len(role_skills))) * 100)
        except Exception:
            target_role = None

    skill_groups = {
        "frontend": ["html", "css", "javascript", "react", "typescript"],
        "backend": ["python", "java", "node", "node.js", "flask", "django", "express"],
        "data_ai": ["sql", "pandas", "numpy", "tensorflow", "pytorch", "machine learning", "ai/ml"],
        "cloud_devops": ["aws", "azure", "gcp", "docker", "kubernetes", "devops", "ci/cd", "linux"],
    }
    category_counts = {
        group: sum(1 for s in skills if s in set(vals)) for group, vals in skill_groups.items()
    }

    suggestions = [
        "Rewrite project bullets using: action + tech stack + measurable impact.",
        "Add one advanced project with deployment link and architecture notes.",
        "Include 2-3 certifications aligned to your target role.",
        "List tools and frameworks in a dedicated technical skills block.",
        "Quantify outcomes (latency reduced, accuracy improved, users impacted).",
    ]

    priorities = [
        {"priority": "P1", "title": "Strengthen project impact bullets", "why": "Recruiters prioritize measurable outcomes."},
        {"priority": "P1", "title": "Close role-critical skill gaps", "why": "Role fit increases when missing core skills are added."},
        {"priority": "P2", "title": "Improve section completeness", "why": "Clear structure improves screening speed and readability."},
        {"priority": "P2", "title": "Add certifications and links", "why": "External proof improves trust and shortlisting rates."},
    ]

    word_buckets = Counter(re.findall(r"\b[a-zA-Z][a-zA-Z\-/+.#]{1,20}\b", text_lower))
    top_terms = [w for w, _ in word_buckets.most_common(12) if w not in {"with", "from", "that", "this", "have", "using"}]

    return {
        "resume_score": final_score,
        "quality_band": quality_band,
        "summary": f"Resume shows {quality_band.lower()} readiness with {word_count} words and {len(skills)} detected skills.",
        "word_count": word_count,
        "strengths": strengths,
        "weaknesses": weaknesses,
        "suggestions": suggestions,
        "extracted_skills": skills,
        "section_presence": sections,
        "metrics": {
            "skill_depth": skill_depth,
            "detail_depth": detail_depth,
            "structure_score": structure_score,
            "action_verb_score": action_verb_score,
        },
        "target_role_insight": {
            "predicted_role": target_role,
            "fit_percentage": fit,
            "missing_skills": missing_skills,
            "role_key_skills": role_skills[:10],
        },
        "skill_category_counts": category_counts,
        "priority_plan": priorities,
        "keyword_focus": top_terms[:8],
    }


# ---------------------------------------------------------
# SERP API COURSE FETCHER
# ---------------------------------------------------------
def get_dynamic_courses(skill):
    try:
        print(f"\n🔍 Fetching courses for: {skill}")

        query = f"best free {skill} course tutorial online 2025"

        search = GoogleSearch({
            "q": query,
            "api_key": SERP_API_KEY,
            "location": "India",
            "hl": "en"
        })

        results = search.get_dict()
        print("📥 SERPAPI keys:", results.keys())

        links = []
        organic = results.get("organic_results")

        if organic:
            for r in organic[:5]:
                if "link" in r:
                    links.append(r["link"])

        print("➡ Extracted links:", links)
        return links

    except Exception as e:
        print("❌ SERPAPI Error:", e)
        return []


def _default_course_links(skill: str) -> List[str]:
    query = skill.strip().replace(" ", "+")
    return [
        f"https://www.youtube.com/results?search_query={query}+full+course",
        f"https://www.coursera.org/search?query={query}",
        f"https://www.freecodecamp.org/news/search/?query={query}",
    ]


def _course_links_for_skill(skill: str) -> List[str]:
    links = get_dynamic_courses(skill)
    cleaned = []
    for link in links:
        value = str(link).strip()
        if value and value not in cleaned:
            cleaned.append(value)
    if cleaned:
        return cleaned
    return _default_course_links(skill)


# ---------------------------------------------------------
# ROADMAP GENERATOR
# ---------------------------------------------------------
def generate_learning_roadmap(missing_skills):
    roadmap = {"7_day": [], "30_day": [], "60_day": []}

    for skill in missing_skills:
        roadmap["7_day"].append({
            "skill": skill,
            "plan": [
                f"Day 1–2: Learn {skill} basics.",
                f"Day 3: Understand core concepts.",
                f"Day 4–5: Practice exercises.",
                f"Day 6: Build 1 mini project.",
                f"Day 7: Revision."
            ]
        })

        roadmap["30_day"].append({
            "skill": skill,
            "plan": [
                f"Week 1: Fundamentals.",
                f"Week 2: Intermediate topics.",
                f"Week 3: Build 2 projects.",
                f"Week 4: Advanced + deploy."
            ]
        })

        roadmap["60_day"].append({
            "skill": skill,
            "plan": [
                f"Month 1: Basics → Intermediate.",
                f"Month 1 end: 2 real projects.",
                f"Month 2: Advanced concepts.",
                f"Month 2 end: Capstone project."
            ]
        })

    return roadmap


def _roadmap_topic_catalog(skill: str) -> List[Dict[str, str]]:
    s = skill.strip().lower()
    topic_map = {
        "c++": [
            {"topic": "Data types, variables, constants", "practice": "Write variable declaration and type-casting snippets."},
            {"topic": "Operators and expressions", "practice": "Solve arithmetic/relational operator exercises."},
            {"topic": "Conditional statements (if, switch)", "practice": "Build decision-based mini programs."},
            {"topic": "Loops (for, while, do-while)", "practice": "Solve pattern and iteration problems."},
            {"topic": "Functions and parameter passing", "practice": "Implement reusable utility functions."},
            {"topic": "Arrays and strings", "practice": "Solve searching and manipulation tasks."},
            {"topic": "Pointers and references", "practice": "Trace pointer memory behavior examples."},
            {"topic": "Object-Oriented Programming", "practice": "Create classes with constructors and inheritance."},
            {"topic": "STL containers and iterators", "practice": "Use vector, map, set in coding tasks."},
            {"topic": "Algorithms and complexity", "practice": "Implement sorting/searching with Big-O comparison."},
            {"topic": "File handling and exception safety", "practice": "Build file read/write exercises with error handling."},
            {"topic": "Advanced C++ (templates, smart pointers)", "practice": "Refactor old pointer code using RAII."},
        ],
        "html": [
            {"topic": "HTML document structure, tags, attributes", "practice": "Build a semantic profile page layout."},
            {"topic": "Headings, paragraphs, links, lists", "practice": "Create a content page with nav and nested lists."},
            {"topic": "Images, media, and embedding", "practice": "Build a media-rich article section."},
            {"topic": "Tables and tabular semantics", "practice": "Create an accessible data table with captions."},
            {"topic": "Forms and input types", "practice": "Build registration form with validation attributes."},
            {"topic": "Semantic layout elements", "practice": "Refactor div-heavy page into semantic sections."},
            {"topic": "Accessibility essentials (ARIA, labels, alt)", "practice": "Audit and fix accessibility issues in sample markup."},
            {"topic": "SEO and metadata", "practice": "Add title/meta/open-graph tags for a landing page."},
        ],
        "css": [
            {"topic": "Selectors, cascade, specificity", "practice": "Resolve selector conflicts in a style challenge."},
            {"topic": "Box model, margin, padding, border", "practice": "Recreate card components with precise spacing."},
            {"topic": "Typography and color systems", "practice": "Build typographic scale and reusable color tokens."},
            {"topic": "Flexbox layout", "practice": "Create responsive navbar and card alignment systems."},
            {"topic": "CSS Grid layout", "practice": "Build dashboard grid with fixed + fluid sections."},
            {"topic": "Responsive design + media queries", "practice": "Make desktop design mobile-first."},
            {"topic": "Transitions, transforms, animations", "practice": "Add interactive micro-animations to components."},
            {"topic": "Architecture and maintainability", "practice": "Refactor styles using utility/component patterns."},
        ],
        "javascript": [
            {"topic": "Syntax, variables, operators, coercion", "practice": "Solve operator and type coercion drills."},
            {"topic": "Conditionals, loops, and control flow", "practice": "Implement branching and iteration challenges."},
            {"topic": "Functions, scope, closures", "practice": "Write closure-based counter and memoization examples."},
            {"topic": "Arrays, objects, destructuring", "practice": "Solve data transformation problems."},
            {"topic": "DOM manipulation and events", "practice": "Build interactive to-do app behavior."},
            {"topic": "Async JS: callbacks, promises, async/await", "practice": "Fetch API data with loading/error states."},
            {"topic": "Error handling and debugging", "practice": "Debug broken scripts and write guarded code."},
            {"topic": "ES6+ patterns and modules", "practice": "Split app into module files and utilities."},
        ],
        "python": [
            {"topic": "Syntax, variables, data types", "practice": "Write beginner scripts using core types."},
            {"topic": "Conditionals, loops, comprehensions", "practice": "Solve logic tasks with clean control flow."},
            {"topic": "Functions and modules", "practice": "Package utilities into reusable modules."},
            {"topic": "Lists, dicts, sets, tuples", "practice": "Implement transformation and lookup problems."},
            {"topic": "File handling and exceptions", "practice": "Parse files and handle failures gracefully."},
            {"topic": "OOP fundamentals", "practice": "Model domain objects with classes and methods."},
            {"topic": "Testing and debugging", "practice": "Write unit tests and fix failing cases."},
            {"topic": "APIs and automation scripts", "practice": "Build script that calls external APIs."},
        ],
        "sql": [
            {"topic": "SELECT, filtering, sorting", "practice": "Query datasets using WHERE and ORDER BY."},
            {"topic": "Aggregations and GROUP BY", "practice": "Write KPI queries with HAVING."},
            {"topic": "JOINS across tables", "practice": "Solve multi-table reporting scenarios."},
            {"topic": "Subqueries and CTEs", "practice": "Refactor nested queries with CTEs."},
            {"topic": "Window functions", "practice": "Implement ranking and running-total analytics."},
            {"topic": "Indexes and optimization", "practice": "Compare query plans pre/post indexing."},
            {"topic": "Transactions and constraints", "practice": "Implement safe insert/update workflows."},
            {"topic": "Advanced case studies", "practice": "Solve interview-grade SQL business problems."},
        ],
    }

    default_topics = [
        {"topic": f"{skill} fundamentals and environment setup", "practice": "Set up tools and run baseline examples."},
        {"topic": f"Core {skill} syntax and constructs", "practice": "Build short programs covering core syntax."},
        {"topic": f"Control flow and data handling in {skill}", "practice": "Solve beginner-to-intermediate practice tasks."},
        {"topic": f"Intermediate {skill} patterns", "practice": "Implement use-cases with reusable abstractions."},
        {"topic": f"Applied {skill} projects", "practice": "Build mini-projects reflecting real requirements."},
        {"topic": f"Advanced {skill} optimization and interview prep", "practice": "Solve timed challenges and review weak areas."},
    ]
    return topic_map.get(s, default_topics)


def _generate_custom_roadmap(skill: str, days: int) -> List[Dict[str, Any]]:
    days = max(3, min(120, int(days)))
    topics = _roadmap_topic_catalog(skill)
    roadmap: List[Dict[str, Any]] = []

    for day in range(1, days + 1):
        progress = day / days
        topic_idx = int(((day - 1) * len(topics)) / days)
        topic_idx = max(0, min(topic_idx, len(topics) - 1))
        selected = topics[topic_idx]
        topic = selected["topic"]
        practice_hint = selected["practice"]

        if progress <= 0.25:
            phase = "Foundation"
            tasks = [
                f"Study: {topic}",
                "Create concise notes with key rules and syntax patterns",
                f"Hands-on: {practice_hint}",
            ]
        elif progress <= 0.6:
            phase = "Core Build"
            tasks = [
                f"Deep dive: {topic}",
                "Solve 4 focused exercises covering normal and edge cases",
                f"Hands-on: {practice_hint}",
            ]
        elif progress <= 0.85:
            phase = "Applied Practice"
            tasks = [
                f"Apply in mini-project module: {topic}",
                f"Hands-on: {practice_hint}",
                "Write one short reflection on mistakes and fixes",
            ]
        else:
            phase = "Advanced + Review"
            tasks = [
                f"Advanced revision: {topic}",
                "Attempt 5 interview-level questions under time limit",
                "Review wrong answers and update cheat sheet",
            ]

        roadmap.append({
            "day": day,
            "phase": phase,
            "focus": topic,
            "tasks": tasks,
        })

    return roadmap


# ---------------------------------------------------------
# 1) SKILL EXTRACTION ENDPOINT
# ---------------------------------------------------------
@app.route("/extract_skills", methods=["POST"])
def extract_skills():
    data = request.json
    text = data.get("resume_text", "")
    extracted = _extract_skills_from_text(text)
    return jsonify({"skills": extracted})


@app.route("/extract_resume_pdf", methods=["POST"])
def extract_resume_pdf():
    if "resume" not in request.files:
        return jsonify({"error": "Missing PDF file with field name 'resume'"}), 400

    file = request.files["resume"]
    if not file or not file.filename:
        return jsonify({"error": "Invalid file"}), 400

    if not file.filename.lower().endswith(".pdf"):
        return jsonify({"error": "Only PDF files are supported"}), 400

    if PdfReader is None:
        return jsonify({"error": "pypdf is not installed in ai-service environment"}), 500

    try:
        reader = PdfReader(file.stream)
        pages = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")
        raw_text = "\n".join(pages)
        resume_text = _normalize_whitespace(raw_text)
        if not resume_text:
            return jsonify({"error": "Could not extract text from PDF"}), 400
        return jsonify({
            "resume_text": resume_text,
            "page_count": len(reader.pages),
            "word_count": len(resume_text.split()),
        })
    except Exception as e:
        return jsonify({"error": f"PDF extraction failed: {str(e)}"}), 500


# ---------------------------------------------------------
# 2) RESUME SCORING ENDPOINT
# ---------------------------------------------------------
@app.route("/resume_score", methods=["POST"])
def resume_score():
    data = request.json
    skills = data.get("skills", [])
    text = data.get("resume_text", "")

    skill_score = min(len(skills) * 5, 50)
    length_score = min(len(text.split()) // 20, 30)

    keywords = ["project", "experience", "developed", "implemented", "built"]
    keyword_score = sum([1 for k in keywords if k in text.lower()])
    keyword_score = min(keyword_score * 4, 20)

    final_score = skill_score + length_score + keyword_score

    strengths = []
    weaknesses = []

    if len(skills) > 5:
        strengths.append("Good skill variety.")
    else:
        weaknesses.append("Too few technical skills.")

    if length_score > 20:
        strengths.append("Good resume detail.")
    else:
        weaknesses.append("Resume lacks detail.")

    if keyword_score > 8:
        strengths.append("Uses strong action words.")
    else:
        weaknesses.append("Not enough action verbs.")

    suggestions = [
        "Add more detailed project descriptions.",
        "Use action words: built, developed, implemented.",
        "Add certifications.",
        "Include quantifiable achievements."
    ]

    return jsonify({
        "resume_score": final_score,
        "strengths": strengths,
        "weaknesses": weaknesses,
        "suggestions": suggestions
    })


@app.route("/analyze_resume_detailed", methods=["POST"])
def analyze_resume_detailed():
    data = request.json or {}
    text = _normalize_whitespace(data.get("resume_text", ""))
    if not text:
        return jsonify({"error": "resume_text is required"}), 400

    skills = _extract_skills_from_text(text)
    analysis = _calculate_resume_analysis(text, skills)
    return jsonify(analysis)


# ---------------------------------------------------------
# 3) CAREER PREDICTION + LINKS + ROADMAP
# ---------------------------------------------------------
@app.route("/predict_role", methods=["POST"])
def predict():
    data = request.json
    user_skills = data.get("skills", [])

    user_skills = [s for s in user_skills if s in encoder.classes_]

    vector = encoder.transform([user_skills])
    pred_role = model.predict(vector)[0]

    row = df[df["job_role"] == pred_role].iloc[0]
    role_skills = row["skills"].split(",")

    matched = list(set(role_skills) & set(user_skills))
    missing = list(set(role_skills) - set(user_skills))

    fit = int((len(matched) / len(role_skills)) * 100)

    learning_links = {s: _course_links_for_skill(s) for s in missing}
    roadmap = generate_learning_roadmap(missing)

    return jsonify({
        "predicted_role": str(pred_role),
        "fit_percentage": fit,
        "matched_skills": matched,
        "missing_skills": missing,
        "recommended_certifications": row["certifications"],
        "salary_range": f"{int(row['min_salary'])} - {int(row['max_salary'])} {row['currency']}",
        "experience_required": int(row["required_experience_years"]),
        "job_level": row["job_level"],
        "location": row["location"],
        "learning_path": learning_links,
        "roadmap": roadmap
    })


@app.route("/learning_courses", methods=["POST", "OPTIONS"])
def learning_courses():
    if request.method == "OPTIONS":
        return jsonify({"ok": True}), 200

    data = request.json or {}
    skills = data.get("skills", [])
    if not isinstance(skills, list) or not skills:
        return jsonify({"error": "skills array is required"}), 400

    cleaned_skills = []
    for skill in skills:
        s = str(skill).strip().lower()
        if s and s not in cleaned_skills:
            cleaned_skills.append(s)

    learning_links = {s: _course_links_for_skill(s) for s in cleaned_skills[:15]}
    return jsonify({"learning_path": learning_links})


@app.route("/generate_skill_roadmap", methods=["POST", "OPTIONS"])
def generate_skill_roadmap():
    if request.method == "OPTIONS":
        return jsonify({"ok": True}), 200

    data = request.json or {}
    skill = str(data.get("skill", "")).strip().lower()
    duration_days = data.get("duration_days", 30)

    if not skill:
        return jsonify({"error": "skill is required"}), 400

    try:
        duration_days = int(duration_days)
    except Exception:
        return jsonify({"error": "duration_days must be a number"}), 400

    if duration_days < 3 or duration_days > 120:
        return jsonify({"error": "duration_days must be between 3 and 120"}), 400

    roadmap = _generate_custom_roadmap(skill, duration_days)
    return jsonify({
        "skill": skill,
        "duration_days": duration_days,
        "roadmap": roadmap,
    })


def _split_lines(value: Any) -> List[str]:
    if isinstance(value, list):
        return [str(v).strip() for v in value if str(v).strip()]
    text = str(value or "").strip()
    if not text:
        return []
    return [line.strip("- ").strip() for line in re.split(r"[\n;]+", text) if line.strip()]


@app.route("/build_resume_docx", methods=["POST", "OPTIONS"])
def build_resume_docx():
    if request.method == "OPTIONS":
        return jsonify({"ok": True}), 200

    if Document is None:
        return jsonify({"error": "python-docx is not installed in ai-service environment"}), 500

    data = request.json or {}
    name = str(data.get("name", "")).strip()
    email = str(data.get("email", "")).strip()
    phone = str(data.get("phone", "")).strip()
    location = str(data.get("location", "")).strip()
    linkedin = str(data.get("linkedin", "")).strip()
    github = str(data.get("github", "")).strip()
    target_role = str(data.get("target_role", "")).strip()
    target_company = str(data.get("target_company", "")).strip()
    summary = str(data.get("summary", "")).strip()
    skills = _split_lines(data.get("skills", []))
    projects = _split_lines(data.get("projects", ""))
    experience = _split_lines(data.get("experience", ""))
    education = _split_lines(data.get("education", ""))
    certifications = _split_lines(data.get("certifications", ""))
    assessments = data.get("passed_assessments", [])

    if not name:
        return jsonify({"error": "name is required"}), 400
    if not target_role:
        return jsonify({"error": "target_role is required"}), 400

    # ATS-friendly structure: simple headings, bullets, and keyword-rich objective.
    doc = Document()
    doc.add_heading(name, level=0)

    contact_parts = [p for p in [phone, email, location] if p]
    links = [p for p in [linkedin, github] if p]
    if contact_parts:
        doc.add_paragraph(" | ".join(contact_parts))
    if links:
        doc.add_paragraph(" | ".join(links))

    doc.add_heading("Professional Summary", level=1)
    if summary:
        doc.add_paragraph(summary)
    else:
        company_text = f" at {target_company}" if target_company else ""
        doc.add_paragraph(
            f"Aspiring {target_role}{company_text} candidate with hands-on technical project exposure and strong problem-solving focus."
        )

    doc.add_heading("Target Role", level=1)
    role_line = f"Applying for {target_role}"
    if target_company:
        role_line += f" at {target_company}"
    doc.add_paragraph(role_line)

    doc.add_heading("Technical Skills", level=1)
    if skills:
        doc.add_paragraph(", ".join(sorted(set(skills))))
    else:
        doc.add_paragraph("Add your core technical skills here.")

    doc.add_heading("Projects", level=1)
    if projects:
        for item in projects:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("Add project title + stack + measurable impact.", style="List Bullet")

    doc.add_heading("Experience", level=1)
    if experience:
        for item in experience:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("Fresher or internship experience can be added here.", style="List Bullet")

    doc.add_heading("Education", level=1)
    if education:
        for item in education:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("Add degree, institution, and graduation year.", style="List Bullet")

    doc.add_heading("Certifications", level=1)
    if certifications:
        for item in certifications:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("Add relevant certifications.", style="List Bullet")

    doc.add_heading("Skill Assessments", level=1)
    if isinstance(assessments, list) and assessments:
        for a in assessments:
            skill = str(a.get("skill", "")).strip()
            pct = int(a.get("percentage", 0))
            if skill:
                doc.add_paragraph(f"{skill.upper()} Skill Test: Passed ({pct}%)", style="List Bullet")
    else:
        doc.add_paragraph("No passed assessments yet.", style="List Bullet")

    doc.add_heading("Keywords for ATS", level=1)
    ats_keywords = sorted(set(skills + [target_role.lower(), target_company.lower() if target_company else ""]))
    ats_keywords = [k for k in ats_keywords if k]
    doc.add_paragraph(", ".join(ats_keywords[:30]) if ats_keywords else target_role)

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    safe_name = re.sub(r"[^a-zA-Z0-9_-]+", "_", name.lower()).strip("_") or "candidate"
    filename = f"{safe_name}_ats_resume_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"

    return send_file(
        stream,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ---------------------------------------------------------
# 4) FIXED LOCAL OLLAMA CHATBOT (gemma:2b)
# ---------------------------------------------------------
@app.route("/chatbot", methods=["POST"])
def chatbot():
    data = request.json
    message = data.get("message", "")
    skill_context = data.get("skills", [])

    try:
        prompt = (
            f"You are a helpful AI tutor.\n"
            f"User skills: {', '.join(skill_context)}\n"
            f"User question: {message}\n"
            f"Give a short, clear explanation.\n"
        )

        response = requests.post(
            "http://localhost:11434/api/generate",
            json={"model": "gemma:2b", "prompt": prompt},
            stream=True
        )

        reply_text = ""

        for line in response.iter_lines():
            if line:
                try:
                    obj = json.loads(line.decode("utf-8"))
                    if "response" in obj:
                        reply_text += obj["response"]
                except Exception:
                    continue

        if not reply_text.strip():
            return jsonify({"response": "⚠ I could not generate a response."})

        return jsonify({"response": reply_text})

    except Exception as e:
        print("❌ Ollama Chatbot Error:", e)
        return jsonify({"response": "⚠ AI Tutor Error. Try again."})
    
# ---------------------------------------------------------
# 5) GENAI TEST GENERATOR (MCQs + CODING - Dynamic)
# ---------------------------------------------------------
def _extract_json_object(raw_text: str) -> Dict[str, Any]:
    text = raw_text.strip()
    # Remove markdown fences if present.
    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```$", "", text)

    # First try full parse.
    try:
        parsed = json.loads(text)
        if isinstance(parsed, dict):
            return parsed
    except Exception:
        pass

    # Fallback: extract first JSON object block.
    match = re.search(r"\{[\s\S]*\}", text)
    if not match:
        raise ValueError("No JSON object found in AI output")

    parsed = json.loads(match.group(0))
    if not isinstance(parsed, dict):
        raise ValueError("Top-level JSON is not an object")
    return parsed


def _normalize_test_payload(test_json: Dict[str, Any], skill: str) -> Dict[str, Any]:
    mcqs: List[Dict[str, Any]] = test_json.get("mcqs", [])
    if not isinstance(mcqs, list):
        mcqs = []

    normalized_mcqs = []
    for i, q in enumerate(mcqs, start=1):
        if not isinstance(q, dict):
            continue
        question = str(q.get("question", "")).strip()
        options = q.get("options", [])
        answer = str(q.get("answer", "")).strip()
        topic = str(q.get("topic", "General")).strip() or "General"
        difficulty = str(q.get("difficulty", "medium")).strip().lower()
        if difficulty not in {"easy", "medium", "hard"}:
            difficulty = "medium"
        if not question or not isinstance(options, list) or len(options) < 2:
            continue
        options = [str(opt).strip() for opt in options if str(opt).strip()]
        if len(options) < 4:
            continue
        if answer not in options:
            answer = options[0]
        normalized_mcqs.append({
            "id": i,
            "question": question,
            "options": options[:4],
            "answer": answer,
            "topic": topic,
            "difficulty": difficulty
        })

    # Ensure minimum viable test so frontend always renders.
    if not normalized_mcqs:
        normalized_mcqs = _build_fallback_mcqs(skill)
    else:
        filtered = [q for q in normalized_mcqs if not _is_generic_question(q["question"])]
        # If many generated questions are generic/non-technical, use curated fallback.
        if len(filtered) < 14:
            normalized_mcqs = _build_fallback_mcqs(skill)
        else:
            normalized_mcqs = filtered[:20]
            while len(normalized_mcqs) < 20:
                normalized_mcqs.append(_build_fallback_mcqs(skill)[len(normalized_mcqs)])

    normalized_mcqs = _shuffle_mcq_options(normalized_mcqs)
    for i, q in enumerate(normalized_mcqs, start=1):
        q["id"] = i

    return {
        "duration_seconds": int(test_json.get("duration_seconds", 1800)),
        "mcqs": normalized_mcqs
    }


def _is_generic_question(question: str) -> bool:
    q = question.lower()
    generic_signals = [
        "manage this situation",
        "team conflict",
        "leadership",
        "stakeholder",
        "communication style",
        "workplace",
        "soft skill",
        "project deadline pressure",
        "behavioral",
    ]
    technical_signals = [
        "<", "/>", "tag", "attribute", "selector", "query", "function", "algorithm",
        "complexity", "api", "schema", "index", "loop", "snippet", "code",
    ]
    if any(token in q for token in technical_signals):
        return False
    return any(token in q for token in generic_signals)


def _shuffle_mcq_options(mcqs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    rng = random.SystemRandom()
    shuffled = []
    for q in mcqs:
        options = list(q.get("options", []))
        answer = q.get("answer", "")
        if len(options) < 2:
            shuffled.append(q)
            continue
        rng.shuffle(options)
        if answer not in options:
            answer = options[0]
        updated = dict(q)
        updated["options"] = options
        updated["answer"] = answer
        shuffled.append(updated)
    return shuffled


def _build_fallback_mcqs(skill: str) -> List[Dict[str, Any]]:
    s = skill.strip().lower()
    if s == "html":
        return _build_html_fallback_mcqs()

    return [
        {"id": 1, "topic": "Debugging", "difficulty": "easy", "question": f"Which practice improves {skill} debugging quality first?", "options": ["Reproduce the issue consistently", "Change multiple modules at once", "Skip logging", "Ignore edge cases"], "answer": "Reproduce the issue consistently"},
        {"id": 2, "topic": "Fundamentals", "difficulty": "easy", "question": f"What is a core beginner checkpoint in {skill}?", "options": ["Understand syntax and control flow", "Premature micro-optimization", "Avoid testing", "Skip documentation"], "answer": "Understand syntax and control flow"},
        {"id": 3, "topic": "Workflow", "difficulty": "easy", "question": f"In {skill}, why use version control for exercises?", "options": ["Track mistakes and improvements", "Reduce readability", "Disable rollback", "Hide progress"], "answer": "Track mistakes and improvements"},
        {"id": 4, "topic": "Practice Strategy", "difficulty": "easy", "question": f"Which habit best supports learning {skill} fundamentals?", "options": ["Daily small problem solving", "Only watching tutorials", "Copy-paste without understanding", "Ignoring errors"], "answer": "Daily small problem solving"},
        {"id": 5, "topic": "Validation", "difficulty": "easy", "question": f"First action when solution output is wrong in {skill}?", "options": ["Check assumptions and sample inputs", "Rewrite everything", "Ignore constraints", "Increase hardware"], "answer": "Check assumptions and sample inputs"},
        {"id": 6, "topic": "Communication", "difficulty": "easy", "question": f"What indicates solid beginner progress in {skill}?", "options": ["Can explain chosen approach clearly", "Can memorize answers only", "Avoids constraints", "Skips validation"], "answer": "Can explain chosen approach clearly"},
        {"id": 7, "topic": "Performance", "difficulty": "medium", "question": f"For performance issues in {skill}, what should you inspect first?", "options": ["Time complexity and hotspots", "UI color theme", "File naming only", "Comment count"], "answer": "Time complexity and hotspots"},
        {"id": 8, "topic": "Testing", "difficulty": "medium", "question": f"Best way to validate a medium-level {skill} solution?", "options": ["Test normal, boundary, and invalid cases", "Test one happy path only", "Skip negative cases", "Check style only"], "answer": "Test normal, boundary, and invalid cases"},
        {"id": 9, "topic": "Refactoring", "difficulty": "medium", "question": f"When refactoring {skill} logic, what reduces regression risk?", "options": ["Add tests before structural changes", "Change code without tests", "Rename variables only", "Increase nesting depth"], "answer": "Add tests before structural changes"},
        {"id": 10, "topic": "Optimization", "difficulty": "medium", "question": f"In {skill}, why compare two alternative approaches?", "options": ["Choose based on constraints and complexity", "Pick longer code always", "Avoid benchmarks", "Ignore memory usage"], "answer": "Choose based on constraints and complexity"},
        {"id": 11, "topic": "Design Tradeoffs", "difficulty": "medium", "question": f"Good intermediate signal in {skill} assessments?", "options": ["Can justify tradeoffs under constraints", "Uses random approach", "Avoids explanation", "Depends on brute force always"], "answer": "Can justify tradeoffs under constraints"},
        {"id": 12, "topic": "Edge Cases", "difficulty": "medium", "question": f"What is a common medium-level pitfall in {skill}?", "options": ["Ignoring edge cases and input limits", "Writing helper functions", "Using meaningful names", "Reviewing complexity"], "answer": "Ignoring edge cases and input limits"},
        {"id": 13, "topic": "Problem Modeling", "difficulty": "hard", "question": f"For hard {skill} problems, what strategy is strongest?", "options": ["Model constraints before coding", "Start coding without plan", "Optimize syntax first", "Skip proof of correctness"], "answer": "Model constraints before coding"},
        {"id": 14, "topic": "Complexity Analysis", "difficulty": "hard", "question": f"How to verify a hard {skill} algorithm is scalable?", "options": ["Analyze worst-case time and memory", "Run one tiny test", "Rely on intuition only", "Ignore asymptotics"], "answer": "Analyze worst-case time and memory"},
        {"id": 15, "topic": "Advanced Quality", "difficulty": "hard", "question": f"What separates advanced {skill} solutions in interviews?", "options": ["Correctness plus optimized complexity", "Only short code", "Only syntactic tricks", "No explanation"], "answer": "Correctness plus optimized complexity"},
        {"id": 16, "topic": "Stress Testing", "difficulty": "hard", "question": f"If two hard {skill} solutions pass samples, best next step?", "options": ["Stress test adversarial inputs", "Submit immediately always", "Delete helper checks", "Ignore memory profile"], "answer": "Stress test adversarial inputs"},
        {"id": 17, "topic": "Recovery Strategy", "difficulty": "hard", "question": f"What is the best fallback when hard {skill} optimization fails?", "options": ["Provide correct baseline and discuss upgrades", "Return empty output", "Skip explanation", "Hide complexity limits"], "answer": "Provide correct baseline and discuss upgrades"},
        {"id": 18, "topic": "Interview Communication", "difficulty": "hard", "question": f"In advanced {skill}, which behavior shows mastery?", "options": ["Clear reasoning, proofs, and tradeoffs", "Random trial-and-error only", "Avoiding constraints discussion", "Copying templates blindly"], "answer": "Clear reasoning, proofs, and tradeoffs"},
    ]


def _build_html_fallback_mcqs() -> List[Dict[str, Any]]:
    return [
        {"id": 1, "topic": "HTML Tags", "difficulty": "easy", "question": "What is the semantic purpose of the <section> tag?", "options": ["It groups related thematic content", "It makes text bold", "It creates a database table", "It adds JavaScript events automatically"], "answer": "It groups related thematic content"},
        {"id": 2, "topic": "Forms", "difficulty": "easy", "question": "Which input type is best for email validation in basic HTML5?", "options": ["type=\"email\"", "type=\"text-email\"", "type=\"mailbox\"", "type=\"string\""], "answer": "type=\"email\""},
        {"id": 3, "topic": "Accessibility", "difficulty": "easy", "question": "Why is the alt attribute important on <img>?", "options": ["Provides text alternatives for screen readers", "Improves CSS specificity", "Executes image compression", "Creates responsive breakpoints"], "answer": "Provides text alternatives for screen readers"},
        {"id": 4, "topic": "Document Structure", "difficulty": "easy", "question": "Which element defines the document metadata container?", "options": ["<head>", "<metahead>", "<header>", "<manifest>"], "answer": "<head>"},
        {"id": 5, "topic": "Lists", "difficulty": "easy", "question": "Which tags are valid for an ordered list with items?", "options": ["<ol> with <li>", "<ul> with <item>", "<list> with <li>", "<ol> with <item>"], "answer": "<ol> with <li>"},
        {"id": 6, "topic": "Links", "difficulty": "easy", "question": "What does target=\"_blank\" do on an anchor tag?", "options": ["Opens link in a new tab/window", "Downloads the linked page", "Adds SEO meta tags", "Blocks cross-origin requests"], "answer": "Opens link in a new tab/window"},
        {"id": 7, "topic": "Semantic Layout", "difficulty": "medium", "question": "For a blog article card component, which semantic structure is best?", "options": ["<article> with <header>, content, and <footer>", "<div> only with inline styles", "<span> for all blocks", "<section> inside <title>"], "answer": "<article> with <header>, content, and <footer>"},
        {"id": 8, "topic": "Code Snippet", "difficulty": "medium", "question": "Given: <label for=\"email\">Email</label><input id=\"email\"> What does the for attribute ensure?", "options": ["Label click focuses the matching input", "Input value is encrypted", "Form auto-submits on focus", "Browser adds placeholder text"], "answer": "Label click focuses the matching input"},
        {"id": 9, "topic": "Tables", "difficulty": "medium", "question": "Which tag should hold table header cells for accessibility?", "options": ["<th>", "<theadcell>", "<tdh>", "<head>"], "answer": "<th>"},
        {"id": 10, "topic": "Use Case", "difficulty": "medium", "question": "You need a navigation area with primary links. Which element is most appropriate?", "options": ["<nav>", "<links>", "<menuitem>", "<route>"], "answer": "<nav>"},
        {"id": 11, "topic": "Embedded Content", "difficulty": "medium", "question": "Which element is preferred for self-hosted video with controls?", "options": ["<video controls>", "<media controls>", "<movie>", "<embedvideo>"], "answer": "<video controls>"},
        {"id": 12, "topic": "Code Snippet", "difficulty": "medium", "question": "In <input name=\"q\" required>, what does required enforce?", "options": ["Field must be filled before form submit", "Input must contain only numbers", "Field becomes read-only", "Input always gets default value"], "answer": "Field must be filled before form submit"},
        {"id": 13, "topic": "Advanced Semantics", "difficulty": "hard", "question": "Which choice best improves both semantics and accessibility for page landmarks?", "options": ["Use <header>, <nav>, <main>, <aside>, <footer> appropriately", "Use only nested <div> blocks", "Put all content in <table>", "Replace headings with styled <span>"], "answer": "Use <header>, <nav>, <main>, <aside>, <footer> appropriately"},
        {"id": 14, "topic": "Code Snippet", "difficulty": "hard", "question": "Snippet: <button aria-expanded=\"false\" aria-controls=\"menu\">Menu</button>. What should update when menu opens?", "options": ["aria-expanded should become true", "aria-controls should be removed", "button must change to <a>", "aria-label must be empty"], "answer": "aria-expanded should become true"},
        {"id": 15, "topic": "SEO + Structure", "difficulty": "hard", "question": "Why is using one meaningful <h1> and structured heading levels recommended?", "options": ["It improves content hierarchy for users and crawlers", "It enables CSS grid automatically", "It minifies the HTML bundle", "It replaces need for meta tags"], "answer": "It improves content hierarchy for users and crawlers"},
        {"id": 16, "topic": "Forms", "difficulty": "hard", "question": "For a terms checkbox that must be accepted, which is correct?", "options": ["<input type=\"checkbox\" required>", "<checkbox required=\"true\">", "<input type=\"accept\">", "<input type=\"checkbox\" validate>"], "answer": "<input type=\"checkbox\" required>"},
        {"id": 17, "topic": "Use Case", "difficulty": "hard", "question": "You need collapsible FAQ content with native semantics. Which tags fit best?", "options": ["<details> and <summary>", "<accordion> and <title>", "<faq> and <item>", "<collapse> and <header>"], "answer": "<details> and <summary>"},
        {"id": 18, "topic": "Code Snippet", "difficulty": "hard", "question": "Given <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\">, what is its primary effect?", "options": ["Improves mobile viewport scaling behavior", "Adds Open Graph tags", "Enables service workers", "Forces desktop rendering on mobile"], "answer": "Improves mobile viewport scaling behavior"},
        {"id": 19, "topic": "Forms", "difficulty": "hard", "question": "Which control is most semantically correct for selecting exactly one value from many choices?", "options": ["A radio button group with same name", "Multiple independent checkboxes", "A text input", "A contenteditable div"], "answer": "A radio button group with same name"},
        {"id": 20, "topic": "Structure", "difficulty": "hard", "question": "Which element should wrap independent, self-contained content that could be syndicated?", "options": ["<article>", "<section>", "<div>", "<aside>"], "answer": "<article>"},
    ]


@app.route("/generate_test", methods=["POST"])
def generate_test():
    data = request.json
    skill = data.get("skill", "programming")
    level = data.get("level", "mixed")
    model_name = os.getenv("OLLAMA_MODEL", "gemma:2b")
    cache_key = f"{skill.strip().lower()}::{level.strip().lower()}::v3"

    # 15-minute cache to avoid regenerating the same test repeatedly.
    cached = TEST_CACHE.get(cache_key)
    now = time.time()
    if cached and (now - cached["created_at"] <= 900):
        return jsonify(cached["payload"])

    try:
        prompt = f"""
You are an expert technical interviewer.

Generate a strict, interview-grade MCQ assessment for skill: {skill}
Difficulty mode: {level}
Target style: HackerRank / HackerEarth screening.

Return ONLY valid JSON in this exact format:

{{
  "duration_seconds": 1800,
  "mcqs": [
    {{
      "id": 1,
      "topic": "Topic name",
      "difficulty": "easy",
      "question": "Question text (concise but challenging)",
      "options": ["A", "B", "C", "D"],
      "answer": "Correct option text"
    }}
  ]
}}

Rules:
- Generate exactly 20 MCQs total.
- Difficulty split must be 7 easy, 7 medium, 6 hard.
- MCQs must be technical and subject-oriented for {skill}.
- At least 7 questions must reference concrete syntax/tags/operators/APIs of {skill}.
- At least 5 questions must include short code snippets and ask output/behavior/correctness.
- At least 4 questions must be applied use-case questions (choose correct construct for a requirement).
- Make questions skill-specific for {skill}, not generic.
- Keep each question and options concise to reduce verbosity.
- Distractors must be plausible and close to the correct answer, not obviously wrong.
- Do NOT create soft-skill/behavioral/workplace-management questions.
- Do not include markdown fences.
- Return ONLY JSON, no explanation.
"""

        response = requests.post(
            "http://localhost:11434/api/generate",
            json={
                "model": model_name,
                "prompt": prompt,
                "stream": True,
                "options": {
                    "temperature": 0.25,
                    "num_predict": 1000
                }
            },
            timeout=(8, 45),
            stream=True
        )
        response.raise_for_status()

        raw_output = ""
        for line in response.iter_lines():
            if not line:
                continue
            try:
                chunk = json.loads(line.decode("utf-8"))
                raw_output += chunk.get("response", "")
                if chunk.get("done"):
                    break
            except Exception:
                continue

        print("Raw Test AI Output preview:\n", str(raw_output)[:500])

        if not raw_output.strip():
            return jsonify({"error": "Empty AI output"}), 500

        test_json = _extract_json_object(raw_output)
        normalized = _normalize_test_payload(test_json, skill)
        TEST_CACHE[cache_key] = {"created_at": now, "payload": normalized}
        return jsonify(normalized)

    except requests.exceptions.RequestException as e:
        print("Ollama request error:", e)
        fallback_mcqs = _shuffle_mcq_options(_build_fallback_mcqs(skill))
        for i, q in enumerate(fallback_mcqs, start=1):
            q["id"] = i
        fallback = {"duration_seconds": 1800, "mcqs": fallback_mcqs, "warning": "GenAI timeout; returned fallback advanced MCQ test."}
        TEST_CACHE[cache_key] = {"created_at": now, "payload": fallback}
        return jsonify(fallback), 200
    except ValueError as e:
        print("Test JSON parsing error:", e)
        return jsonify({"error": f"Invalid AI JSON output: {str(e)}"}), 500
    except Exception as e:
        print("Test Generation Error:", e)
        return jsonify({"error": "AI test generation failed"}), 500
# ---------------------------------------------------------
# HEALTH CHECK
# ---------------------------------------------------------
@app.route("/")
def home():
    return "AI Career Prediction API Running Successfully!"


# ---------------------------------------------------------
# START SERVER
# ---------------------------------------------------------
if __name__ == "__main__":
    app.run(port=8000)



